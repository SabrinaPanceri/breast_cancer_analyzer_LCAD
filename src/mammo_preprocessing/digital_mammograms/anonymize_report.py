#https://python-docx.readthedocs.io/en/latest/
from docx import Document
from docx.shared import Inches
import os
import sys
import glob


def open_docx(file):
	doc = Document(file)
	return doc


def process_document (root_list, report_list, files_list, input_folder, output_folder):
	if not input_folder.endswith(os.sep):
		input_folder += os.sep
	if not output_folder.endswith(os.sep):
		output_folder += os.sep

	for i in range(len(report_list)):
		print(report_list[i] + "\n")
		output_path = root_list[i].replace(input_folder, output_folder)
		os.makedirs(output_path, exist_ok = True)


		doc = open_docx (report_list[i])
		doc.paragraphs[0].text = 'XXXXXX'
		cont = 0
		for p in doc.paragraphs:
			#print p.text
			if 'Paciente:' in p.text:

				break
			cont=cont+1	

		first_two_dots = False
		is_Data = False
		first_two_dots_pos = 0
		Data_pos = 0
		#print len(doc.paragraphs[cont].text)
		print(doc.paragraphs[cont].text + "\n")	
		
		aux = doc.paragraphs[cont].text.split("DatadeNascimento")

		print(aux)

		aux[0] = "Paciente: XXXXXXXXXX\n"
		doc.paragraphs[cont].text = str(aux[0]) + str(aux[1])
		output = os.path.join(output_path, "anon_" + files_list[i])

		doc.save(output)


def get_reports_filename_in_folder(folder):
	report_list = []
	files_list = []
	root_list = []
	for root, dirs, files in os.walk(folder):
		for file in files:
			if file.endswith('.docx'):
				root_list.append(root)
				report_list.append(os.path.join(root, file))
				files_list.append(file)

	return root_list, report_list, files_list


def main():
	#print len(sys.argv)
	if len(sys.argv) < 3:
		print("Incorrect input!")
		print("The right way: python report_editor.py <reports_folder> <output_folder>")
		exit(1)

	report_list = []
	files_list = []
	root_list, report_list, files_list = get_reports_filename_in_folder(sys.argv[1])
	#doc = open_docx()
	process_document (root_list, report_list, files_list, sys.argv[1], sys.argv[2])
  
if __name__== "__main__":
	main()